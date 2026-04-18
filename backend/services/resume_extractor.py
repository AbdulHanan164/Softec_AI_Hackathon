import io
import re
import json
from backend.config import MODEL, TEMPERATURE, client

# ─────────────────────────────────────────────────────────────────────────────
# REGEX EXTRACTORS  (always run, never fail)
# ─────────────────────────────────────────────────────────────────────────────

def _regex_phone(text: str) -> str:
    patterns = [
        r'(\+92[\s\-]?\d{3}[\s\-]?\d{7})',          # Pakistani +92
        r'(0\d{2,3}[\s\-]?\d{7,8})',                 # Local 03xx
        r'(\+\d{1,3}[\s\-]?\(?\d{2,4}\)?[\s\-]?\d{3,4}[\s\-]?\d{4})',  # International
        r'(\d{4}[\s\-]\d{3}[\s\-]\d{4})',            # Pattern: 0300 123 4567
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            num = m.group(1).strip()
            digits = re.sub(r'\D', '', num)
            if len(digits) >= 10:
                return num
    return ''


def _regex_linkedin(text: str) -> str:
    # Most reliable: look for the URL directly in text
    m = re.search(r'linkedin\.com/in/([\w\-]+)', text, re.IGNORECASE)
    if m:
        return f"https://linkedin.com/in/{m.group(1)}"
    m = re.search(r'linkedin\.com/([\w\-/]+)', text, re.IGNORECASE)
    if m:
        return f"https://linkedin.com/{m.group(1).strip('/')}"
    # "LinkedIn: username" pattern
    m = re.search(r'linkedin[:\s]+@?([\w\-]+)', text, re.IGNORECASE)
    if m and '.' not in m.group(1):
        return f"https://linkedin.com/in/{m.group(1)}"
    return ''


def _regex_github(text: str) -> str:
    m = re.search(r'github\.com/([\w\-]+)', text, re.IGNORECASE)
    if m:
        uname = m.group(1)
        if uname.lower() not in ('com', 'io', 'www'):
            return f"https://github.com/{uname}"
    m = re.search(r'github[:\s]+@?([\w\-]+)', text, re.IGNORECASE)
    if m and '.' not in m.group(1):
        uname = m.group(1)
        if uname.lower() not in ('com', 'io', 'www'):
            return f"https://github.com/{uname}"
    return ''


def _regex_email(text: str) -> str:
    m = re.search(r'[\w\.\+\-]+@[\w\-]+\.[a-zA-Z]{2,}', text)
    return m.group(0) if m else ''


def _regex_city(text: str) -> str:
    # Scan first 600 chars (header area)
    header = text[:600]
    # Pattern: "City, Country" near contact info
    m = re.search(
        r'\b(Lahore|Karachi|Islamabad|Rawalpindi|Multan|Faisalabad|Peshawar|Quetta|'
        r'Hyderabad|Sialkot|Gujranwala|Bahawalpur|'
        r'Dubai|Abu Dhabi|Riyadh|Jeddah|Doha|Kuwait|Muscat|'
        r'London|Toronto|New York|Sydney|Kuala Lumpur|Singapore)'
        r'(?:[,\s]+(?:Pakistan|UAE|UK|Canada|USA|KSA|Qatar|Kuwait|Oman|Malaysia|Australia))?',
        header, re.IGNORECASE
    )
    if m:
        return m.group(0).strip()
    # Generic "City, Country" pattern
    m = re.search(r'\b([A-Z][a-z]+),\s*([A-Z][a-z]+)\b', header)
    if m:
        return m.group(0).strip()
    return ''


def _regex_cgpa(text: str) -> float | None:
    patterns = [
        r'(?:CGPA|GPA|Grade\s*Point)[:\s]+(\d+\.?\d*)\s*/\s*(\d+\.?\d*)',
        r'(?:CGPA|GPA)[:\s]+(\d+\.?\d*)',
        r'(\d+\.?\d*)\s*/\s*(4\.0|5\.0)\s*(?:CGPA|GPA|scale)',
    ]
    for p in patterns:
        m = re.search(p, text, re.IGNORECASE)
        if m:
            try:
                val = float(m.group(1))
                if m.lastindex and m.lastindex >= 2:
                    scale = float(m.group(2))
                    if scale > 4.0:
                        val = round((val / scale) * 4.0, 2)
                return max(0.0, min(4.0, val))
            except (ValueError, IndexError):
                continue
    return None


def _regex_name(text: str) -> str:
    """Try to get name from first non-empty line of CV."""
    for line in text.split('\n')[:6]:
        line = line.strip()
        # Name line: 2-4 words, all title case, no digits/symbols
        if line and 2 <= len(line.split()) <= 5 and re.match(r'^[A-Za-z\s\.\-]+$', line):
            return line
    return ''


# ─────────────────────────────────────────────────────────────────────────────
# FILE READERS
# ─────────────────────────────────────────────────────────────────────────────

def extract_text_from_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        pages = [page.extract_text() or '' for page in reader.pages]
        return '\n'.join(pages)
    except Exception as e:
        raise ValueError(f"Could not read PDF: {e}")


def extract_text_from_docx(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also read table cells — contact info is often in header tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return '\n'.join(lines)
    except Exception as e:
        raise ValueError(f"Could not read DOCX: {e}")


def extract_text_from_txt(content: bytes) -> str:
    return content.decode('utf-8', errors='ignore')


# ─────────────────────────────────────────────────────────────────────────────
# AI EXTRACTION PROMPT
# ─────────────────────────────────────────────────────────────────────────────

EXTRACT_PROMPT = """You are a precise resume parser. Extract ALL information from this CV.

Return ONLY a JSON object — no markdown, no code fences, no explanation.

{
  "name": "full name",
  "degree_program": "degree + major e.g. BS Computer Science",
  "semester": <int 1-8>,
  "cgpa": <float 0.0-4.0, convert if on 5.0 scale>,
  "skills": ["every skill, tool, language, framework found"],
  "interests": ["interests from objective, hobbies, activities"],
  "preferred_opportunity_types": ["internship","scholarship","competition","fellowship","admission"],
  "financial_need": "high"|"medium"|"low",
  "location_preference": "remote"|"onsite"|"any",
  "past_experience": ["Title at Company (Year) — one liner per role"],
  "phone": "phone number exactly as written, empty string if not found",
  "linkedin": "linkedin URL or username, empty string if not found",
  "github": "github URL or username, empty string if not found",
  "city": "City, Country if found, empty string if not found"
}

Rules:
- Extract EVERY skill: languages, frameworks, libraries, tools, software
- ALL experience: internships, jobs, projects, research, volunteering
- Semester: 1st year=1-2, 2nd year=3-4, 3rd year=5-6, 4th year=7-8
- CGPA on 5.0 scale: divide by 1.25. On 10.0: divide by 2.5
- Return "" for missing fields, never null
"""


# ─────────────────────────────────────────────────────────────────────────────
# MAIN EXTRACTOR
# ─────────────────────────────────────────────────────────────────────────────

def extract_profile_from_resume(filename: str, content: bytes) -> dict:
    fname = filename.lower()

    if fname.endswith('.pdf'):
        text = extract_text_from_pdf(content)
    elif fname.endswith('.docx') or fname.endswith('.doc'):
        text = extract_text_from_docx(content)
    elif fname.endswith('.txt'):
        text = extract_text_from_txt(content)
    else:
        raise ValueError("Unsupported file type. Use PDF, DOCX, or TXT.")

    if not text.strip():
        raise ValueError("No text could be extracted. Make sure the CV is not a scanned image.")

    print(f"[CV] Extracted {len(text)} chars")
    print(f"[CV] Preview:\n{text[:400]}\n{'─'*40}")

    # ── STEP 1: Regex extraction (always succeeds) ────────────────────────
    r_phone    = _regex_phone(text)
    r_linkedin = _regex_linkedin(text)
    r_github   = _regex_github(text)
    r_city     = _regex_city(text)
    r_cgpa     = _regex_cgpa(text)
    r_name     = _regex_name(text)

    print(f"[CV] Regex → phone={r_phone!r} linkedin={r_linkedin!r} github={r_github!r} city={r_city!r} cgpa={r_cgpa} name={r_name!r}")

    # Start with a solid baseline from regex alone
    profile: dict = {
        'name':            r_name,
        'degree_program':  '',
        'semester':        1,
        'cgpa':            r_cgpa if r_cgpa is not None else 3.0,
        'skills':          [],
        'interests':       [],
        'preferred_opportunity_types': ['internship'],
        'financial_need':      'medium',
        'location_preference': 'any',
        'past_experience': [],
        'phone':    r_phone,
        'linkedin': r_linkedin,
        'github':   r_github,
        'city':     r_city,
        'open_to_hiring': False,
    }

    # ── STEP 2: AI extraction (best-effort, graceful fallback) ────────────
    try:
        response = client.chat.completions.create(
            model=MODEL,
            temperature=0,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": EXTRACT_PROMPT},
                {"role": "user",   "content": f"CV:\n{text[:7000]}"},
            ],
        )
        raw = response.choices[0].message.content or '{}'
        print(f"[CV] AI raw (first 300): {raw[:300]}")

        ai = json.loads(raw)

        # Merge AI into profile — AI wins for text fields, regex wins for contact
        profile['name']           = ai.get('name', '')           or profile['name']
        profile['degree_program'] = ai.get('degree_program', '')  or ''
        profile['financial_need'] = ai.get('financial_need', 'medium') or 'medium'
        profile['location_preference'] = ai.get('location_preference', 'any') or 'any'

        # Semester
        try:
            profile['semester'] = max(1, min(8, int(ai.get('semester', 1))))
        except (TypeError, ValueError):
            profile['semester'] = 1

        # CGPA — regex wins if found, else use AI value
        if r_cgpa is None:
            try:
                ai_cgpa = float(ai.get('cgpa', 3.0))
                profile['cgpa'] = max(0.0, min(4.0, ai_cgpa))
            except (TypeError, ValueError):
                profile['cgpa'] = 3.0

        # Lists — use AI values (more complete)
        for field in ('skills', 'interests', 'past_experience', 'preferred_opportunity_types'):
            val = ai.get(field, [])
            if isinstance(val, list) and len(val) > 0:
                profile[field] = [str(v).strip() for v in val if v]

        # Contact — regex always wins, AI only fills if regex found nothing
        if not profile['phone']    and ai.get('phone'):
            profile['phone']    = str(ai['phone']).strip()
        if not profile['linkedin'] and ai.get('linkedin'):
            profile['linkedin'] = _normalise_linkedin(str(ai['linkedin']))
        if not profile['github']   and ai.get('github'):
            profile['github']   = _normalise_github(str(ai['github']))
        if not profile['city']     and ai.get('city'):
            profile['city']     = str(ai['city']).strip()

        print(f"[CV] AI merge done ✓")

    except Exception as ai_err:
        # AI failed — regex data is still returned, skills/experience will be empty
        print(f"[CV] AI extraction failed: {ai_err} — returning regex data only")

    # ── STEP 3: De-duplicate lists ────────────────────────────────────────
    for field in ('skills', 'interests', 'past_experience'):
        seen, out = set(), []
        for item in profile[field]:
            key = item.strip().lower()
            if key and key not in seen:
                seen.add(key)
                out.append(item.strip())
        profile[field] = out

    print(f"[CV] Final → name={profile['name']!r} phone={profile['phone']!r} "
          f"linkedin={profile['linkedin']!r} github={profile['github']!r} "
          f"city={profile['city']!r} skills={len(profile['skills'])} exp={len(profile['past_experience'])}")

    return profile


# ─────────────────────────────────────────────────────────────────────────────
# URL NORMALIZERS
# ─────────────────────────────────────────────────────────────────────────────

def _normalise_linkedin(val: str) -> str:
    val = val.strip() if val else ''
    if not val:
        return ''
    if 'linkedin.com' in val.lower():
        return val if val.startswith('http') else f'https://{val}'
    slug = val.lstrip('@').strip('/')
    if slug and ' ' not in slug:
        return f'https://linkedin.com/in/{slug}'
    return ''


def _normalise_github(val: str) -> str:
    val = val.strip() if val else ''
    if not val:
        return ''
    if 'github.com' in val.lower():
        return val if val.startswith('http') else f'https://{val}'
    slug = val.lstrip('@').strip('/')
    if slug and ' ' not in slug and slug.lower() not in ('com', 'io', 'www'):
        return f'https://github.com/{slug}'
    return ''
